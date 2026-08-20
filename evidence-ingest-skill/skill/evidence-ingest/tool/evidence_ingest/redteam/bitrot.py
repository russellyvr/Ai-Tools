"""bitrot — integrity-corruption agent.

Traps: post-hash bit flips in the vault (validate must detect the byte
mismatch), truncated PDFs (no %%EOF), appended trailer bytes after the PDF
%%EOF / PNG IEND, and magic/extension mismatches. Oracles: structural traps
are dropped at scan with TAMPER_* codes; the post-hash flip is caught by
validate re-hashing (fail-closed). Controls: well-formed PDF, PNG, and text
files are accepted.
"""
from __future__ import annotations

import json

from evidence_ingest.hashing import sha256_of_bytes
from evidence_ingest.schemas import (
    REASON_MAGIC_MISMATCH, REASON_TRAILING, REASON_TRUNCATED)

MINIMAL_PDF = (
    b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    b"2 0 obj\n<< /Type /Pages /Kids [] /Count 0 >>\nendobj\n"
    b"xref\n0 3\ntrailer\n<< /Size 3 /Root 1 0 R >>\nstartxref\n9\n%%EOF\n"
)
MINIMAL_PNG = (
    b"\x89PNG\r\n\x1a\n"
    b"\x00\x00\x00\x0dIHDR" + b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x00\x00\x00\x00"
    + b"\x3a\x7e\x9b\x55"
    + b"\x00\x00\x00\x00IEND\xae\x42\x60\x82"
)


def _drop_reasons(work) -> dict[str, str]:
    out = {}
    with open(work / "ledger" / "drops.jsonl", encoding="utf-8") as f:
        for line in f:
            if line.strip():
                d = json.loads(line)
                out[d["relpath"]] = d["reason_code"]
    return out


def _fixture_docx() -> bytes:
    """Deterministic-enough per-run docx fixture built with python-docx."""
    import io

    import docx as _docx
    d = _docx.Document()
    d.add_paragraph("Quarterly compliance attestation, fiscal 2024.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "account"
    t.cell(0, 1).text = "balance"
    t.cell(1, 0).text = "operating"
    t.cell(1, 1).text = "1000.00"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _fixture_xlsx() -> bytes:
    import io

    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "ledger"
    ws.append(["account", "balance"])
    ws.append(["operating", 1000.0])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def run(ctx) -> None:
    inp, work = ctx.dirs("main")

    (inp / "good.pdf").write_bytes(MINIMAL_PDF)
    (inp / "good.png").write_bytes(MINIMAL_PNG)
    good_txt = b"Ledger of accounts, fiscal 2024. Balanced and reconciled.\n" * 4
    (inp / "good.txt").write_bytes(good_txt)
    good_docx = _fixture_docx()
    good_xlsx = _fixture_xlsx()
    (inp / "good.docx").write_bytes(good_docx)
    (inp / "good.xlsx").write_bytes(good_xlsx)

    (inp / "truncated.pdf").write_bytes(MINIMAL_PDF[:-8])            # loses %%EOF
    (inp / "trailer.pdf").write_bytes(MINIMAL_PDF + b"EXFILTRATED-PAYLOAD")
    (inp / "trailer.png").write_bytes(MINIMAL_PNG + b"\x00hidden")
    (inp / "notreally.pdf").write_bytes(b"just plain text pretending")  # ext/magic
    (inp / "image.txt").write_bytes(MINIMAL_PNG)                      # magic/ext
    (inp / "truncated.docx").write_bytes(good_docx[:-30])            # loses zip EOCD
    (inp / "notreally.docx").write_bytes(b"plain text posing as OOXML")
    (inp / "sheet.csv").write_bytes(MINIMAL_PNG)                     # binary as csv

    result, _ = ctx.pipeline(inp, work)
    reasons = _drop_reasons(work)
    records = json.loads((work / "ledger" / "records.json").read_text(encoding="utf-8"))

    ctx.trap("pdf-truncation-dropped",
             reasons.get("truncated.pdf") == REASON_TRUNCATED,
             f"got {reasons.get('truncated.pdf')}")
    ctx.trap("pdf-appended-trailer-dropped",
             reasons.get("trailer.pdf") == REASON_TRAILING,
             f"got {reasons.get('trailer.pdf')}")
    ctx.trap("png-appended-trailer-dropped",
             reasons.get("trailer.png") == REASON_TRAILING,
             f"got {reasons.get('trailer.png')}")
    ctx.trap("text-as-pdf-extension-dropped",
             reasons.get("notreally.pdf") == REASON_MAGIC_MISMATCH,
             f"got {reasons.get('notreally.pdf')}")
    ctx.trap("png-as-txt-extension-dropped",
             reasons.get("image.txt") == REASON_MAGIC_MISMATCH,
             f"got {reasons.get('image.txt')}")
    ctx.trap("docx-truncation-dropped",
             reasons.get("truncated.docx") == REASON_TRUNCATED,
             f"got {reasons.get('truncated.docx')}")
    ctx.trap("text-as-docx-extension-dropped",
             reasons.get("notreally.docx") == REASON_MAGIC_MISMATCH,
             f"got {reasons.get('notreally.docx')}")
    ctx.trap("png-as-csv-extension-dropped",
             reasons.get("sheet.csv") == REASON_MAGIC_MISMATCH,
             f"got {reasons.get('sheet.csv')}")

    # Controls must be present before the post-hash tamper trap
    pdf_sha = sha256_of_bytes(MINIMAL_PDF)
    png_sha = sha256_of_bytes(MINIMAL_PNG)
    txt_sha = sha256_of_bytes(good_txt)
    docx_sha = sha256_of_bytes(good_docx)
    xlsx_sha = sha256_of_bytes(good_xlsx)
    ctx.control("wellformed-pdf-accepted", pdf_sha in records)
    ctx.control("wellformed-png-accepted", png_sha in records)
    ctx.control("wellformed-txt-accepted", txt_sha in records)
    ctx.control("wellformed-docx-accepted", docx_sha in records)
    ctx.control("wellformed-xlsx-accepted", xlsx_sha in records)
    # docx/xlsx extraction produced real text through the deterministic parsers
    bundle_docx = json.loads(
        (work / "extracted" / f"{docx_sha}.bundle.json").read_text(encoding="utf-8"))
    ctx.control("docx-extracted-native-text",
                bundle_docx["channel"] == "native-text"
                and any("compliance attestation" in p["text"]
                        for p in bundle_docx["pages"]),
                f"channel={bundle_docx['channel']}")
    bundle_xlsx = json.loads(
        (work / "extracted" / f"{xlsx_sha}.bundle.json").read_text(encoding="utf-8"))
    ctx.control("xlsx-extracted-native-text",
                bundle_xlsx["channel"] == "native-text"
                and any("operating\t1000" in p["text"]
                        for p in bundle_xlsx["pages"]),
                f"channel={bundle_xlsx['channel']}")
    ok, detail = ctx.validate_ok(inp, work)
    ctx.control("clean-corpus-validates", ok, detail)

    # Trap: post-hash single-bit flip in a vault object -> validate MUST fail
    vault_obj = work / records[txt_sha]["vault_relpath"]
    data = bytearray(vault_obj.read_bytes())
    data[7] ^= 0x01
    vault_obj.write_bytes(bytes(data))
    ok, _ = ctx.validate_ok(inp, work)
    ctx.trap("post-hash-bit-flip-detected", not ok,
             "validate passed despite flipped bit in vault object")
    # restore for hygiene
    data[7] ^= 0x01
    vault_obj.write_bytes(bytes(data))
